import os
import pytest
from PIL import Image
import docx

from models.config import FormattingRulesConfig, TitlePageData, ProjectPreset
from models.state import ProjectState
from core.renderer import DocxRenderer

def test_config_and_state():
    """Тест работы конфигурационных моделей и автосохранения состояния проекта"""
    rules = FormattingRulesConfig(font_size_pt=14, margin_left_cm=3.0)
    assert rules.font_name == "Times New Roman"
    assert rules.margin_left_cm == 3.0

    title = TitlePageData(topic="Тестовая курсовая работа по программированию")
    assert title.topic == "Тестовая курсовая работа по программированию"

    preset = ProjectPreset(preset_id="it", preset_name="IT Пресет", formatting=rules)
    assert preset.preset_id == "it"

    state = ProjectState(topic="Разработка приложения", project_type="coursework_it")
    state.add_section_content("1.1", "Текст первой подглавы...")
    
    test_save_path = "test_project.courseproject"
    state.save_to_file(test_save_path)
    assert os.path.exists(test_save_path)

    loaded_state = ProjectState.load_from_file(test_save_path)
    assert loaded_state.topic == "Разработка приложения"
    assert loaded_state.sections_content["1.1"] == "Текст первой подглавы..."

    # Очистка за собой
    if os.path.exists(test_save_path):
        os.remove(test_save_path)

def test_docx_renderer_full(tmp_path):
    """Тест всех методов движка верстки DocxRenderer"""
    config = FormattingRulesConfig()
    renderer = DocxRenderer(config)
    
    doc = renderer.create_document()

    # 1. Титульный лист
    title_data = TitlePageData(
        university="КУБАНСКИЙ ГОСУДАРСТВЕННЫЙ АГРАРНЫЙ УНИВЕРСИТЕТ",
        faculty="Факультет прикладной информатики",
        department="Кафедра компьютерных технологий",
        topic="Разработка платежного терминала на C#",
        student_info="Выполнил: студент Иванов И.И.",
        teacher_info="Проверил: доцент Петров П.П.",
        city="Краснодар",
        year=2026
    )
    renderer.render_title_page(doc, title_data)

    # 2. Оглавление
    plan = [
        {"title": "ВВЕДЕНИЕ", "page": 3, "is_section_header": True},
        {"title": "1. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ", "page": 4, "is_section_header": True},
        {"title": "1.1 Описание предметной области", "page": 4, "is_section_header": False},
        {"title": "2. РАЗРАБОТКА ПРИЛОЖЕНИЯ", "page": 10, "is_section_header": True},
    ]
    renderer.add_table_of_contents_placeholder(doc, plan)

    # 3. Заголовки и параграфы
    renderer.add_heading_1(doc, "1. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ")
    renderer.add_heading_2(doc, "1.1 Описание предметной области")
    renderer.add_paragraph(doc, "Данная курсовая работа посвящена вопросам разработки современных программных систем...")

    # 4. Рамка с кодом
    renderer.add_code_block(doc, "public class PaymentTerminal {\n    public void ProcessPayment() {\n        Console.WriteLine(\"Payment Success\");\n    }\n}")

    # 5. Умное сжатие изображения (Smart Image Fitter)
    test_img_path = str(tmp_path / "test_chart.png")
    img = Image.new('RGB', (800, 600), color=(73, 109, 137))
    img.save(test_img_path)
    
    renderer.insert_smart_image(doc, test_img_path, "Рисунок 1.1 – Схема структуры приложения")

    # 6. Таблица данных
    headers = ["№ п/п", "Наименование показателя", "Значение"]
    rows = [
        ["1", "Количество пользователей", "1500"],
        ["2", "Время отклика системы (мс)", "120"],
    ]
    renderer.render_table(doc, "Таблица 1.1 – Нагрузочные показатели системы", headers, rows)

    # 7. Сохранение и верификация читаемости Word файла
    output_docx = str(tmp_path / "test_output.docx")
    doc.save(output_docx)
    
    assert os.path.exists(output_docx)
    
    # Открываем обратно и проверяем количество парграфов и таблиц
    reopened_doc = docx.Document(output_docx)
    assert len(reopened_doc.paragraphs) > 5
    assert len(reopened_doc.tables) >= 3  # Левая таблица авторов + Код + Таблица данных + Таблица Оглавления
