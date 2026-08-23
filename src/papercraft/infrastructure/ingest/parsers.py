"""Local, provenance-preserving parsers for supported input formats."""

from __future__ import annotations

import ast
import csv
import io
import re
from abc import ABC, abstractmethod
from collections.abc import Iterable
from pathlib import Path
from typing import Any

from papercraft.domain import Source

from ._domain import fragment, locator
from .classification import CODE_SUFFIXES, IMAGE_SUFFIXES
from .types import OptionalDependencyError, ParseResult, UnsupportedSourceError
from .vision import VisionOCRPort


def _path(source: Source) -> Path:
    return Path(source.stored_path)


def _decode(data: bytes) -> tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-16", "cp1251"):
        try:
            return data.decode(encoding), encoding
        except (UnicodeDecodeError, UnicodeError):
            continue
    return data.decode("utf-8", errors="replace"), "utf-8-replace"


def _clean_cell(value: Any, max_length: int = 10_000) -> str:
    if value is None:
        return ""
    rendered = str(value).replace("\x00", "")
    return rendered[:max_length]


def _column_name(index: int) -> str:
    letters = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        letters = chr(65 + remainder) + letters
    return letters or "A"


class SourceParser(ABC):
    suffixes: frozenset[str] = frozenset()

    def supports(self, source: Source) -> bool:
        return _path(source).suffix.casefold() in self.suffixes

    @abstractmethod
    def parse(self, source: Source) -> ParseResult:
        raise NotImplementedError


class TextParser(SourceParser):
    suffixes = frozenset({".txt", ".md", ".markdown", ".rst", ".tex", ".ini", ".cfg"})

    def __init__(self, lines_per_fragment: int = 120) -> None:
        self.lines_per_fragment = max(1, lines_per_fragment)

    def parse(self, source: Source) -> ParseResult:
        path = _path(source)
        text, encoding = _decode(path.read_bytes())
        result = ParseResult(source.id, metadata={"encoding": encoding, "kind": "text"})
        lines = text.splitlines()
        current_section: str | None = None
        for start in range(0, len(lines), self.lines_per_fragment):
            selected = lines[start : start + self.lines_per_fragment]
            if path.suffix.casefold() in {".md", ".markdown"}:
                headings = [line.lstrip("#").strip() for line in selected if re.match(r"^#{1,6}\s", line)]
                if headings:
                    current_section = headings[-1]
            content = "\n".join(selected).strip()
            if not content:
                continue
            start_line = start + 1
            end_line = start + len(selected)
            result.fragments.append(
                fragment(
                    source_id=source.id,
                    content=content,
                    source_locator=locator(
                        source_id=source.id,
                        path=path,
                        line_start=start_line,
                        line_end=end_line,
                        section=current_section,
                    ),
                    metadata={"encoding": encoding},
                    ordinal=f"lines:{start_line}-{end_line}",
                )
            )
        if not result.fragments:
            result.warnings.append("empty-text")
        return result


class CodeParser(SourceParser):
    suffixes = CODE_SUFFIXES | frozenset({".json", ".xml", ".ipynb"})

    def __init__(self, lines_per_fragment: int = 160) -> None:
        self.lines_per_fragment = max(1, lines_per_fragment)

    def parse(self, source: Source) -> ParseResult:
        path = _path(source)
        text, encoding = _decode(path.read_bytes())
        lines = text.splitlines()
        symbols = self._symbols(path, text)
        analysis = self._analysis(path, text, symbols)
        result = ParseResult(
            source.id,
            metadata={
                "encoding": encoding,
                "language": path.suffix.lstrip(".").casefold(),
                "symbols": symbols,
                "code_analysis": analysis,
            },
        )
        for start in range(0, len(lines), self.lines_per_fragment):
            selected = lines[start : start + self.lines_per_fragment]
            content = "\n".join(selected).strip()
            if not content:
                continue
            start_line = start + 1
            end_line = start + len(selected)
            chunk_symbols = [
                symbol for symbol in symbols if start_line <= symbol.get("line", 0) <= end_line
            ]
            result.fragments.append(
                fragment(
                    source_id=source.id,
                    content=content,
                    source_locator=locator(
                        source_id=source.id,
                        path=path,
                        line_start=start_line,
                        line_end=end_line,
                    ),
                    metadata={"symbols": chunk_symbols, "language": path.suffix.lstrip(".")},
                    ordinal=f"code:{start_line}-{end_line}",
                )
            )
        if not result.fragments:
            result.warnings.append("empty-code-file")
        return result

    @staticmethod
    def _symbols(path: Path, text: str) -> list[dict[str, Any]]:
        symbols: list[dict[str, Any]] = []
        if path.suffix.casefold() == ".py":
            try:
                tree = ast.parse(text)
            except SyntaxError as error:
                return [{"kind": "syntax_error", "line": error.lineno or 1, "name": error.msg}]
            for node in ast.walk(tree):
                if isinstance(node, (ast.ClassDef, ast.FunctionDef, ast.AsyncFunctionDef)):
                    symbols.append(
                        {
                            "kind": "class" if isinstance(node, ast.ClassDef) else "function",
                            "name": node.name,
                            "line": node.lineno,
                            "end_line": getattr(node, "end_lineno", node.lineno),
                        }
                    )
        else:
            pattern = re.compile(
                r"^\s*(?:(?:export|public|private|protected|static|async)\s+)*"
                r"(?:class|interface|enum|function|def|fn)\s+([A-Za-z_$][\w$]*)",
                flags=re.MULTILINE,
            )
            for match in pattern.finditer(text):
                symbols.append(
                    {"kind": "symbol", "name": match.group(1), "line": text.count("\n", 0, match.start()) + 1}
                )
        return sorted(symbols, key=lambda item: (item["line"], item["name"]))

    @staticmethod
    def _analysis(path: Path, text: str, symbols: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
        """Portable tree-sitter-shaped output with AST precision for Python.

        The app deliberately keeps grammars optional; deployments with
        tree-sitter can substitute an analyser without changing this contract.
        """

        dependencies: list[dict[str, Any]] = []
        entrypoints: list[dict[str, Any]] = []
        endpoints: list[dict[str, Any]] = []
        tests: list[dict[str, Any]] = []
        if path.suffix.casefold() == ".py":
            try:
                tree = ast.parse(text)
            except SyntaxError:
                tree = None
            if tree is not None:
                for node in ast.walk(tree):
                    if isinstance(node, ast.Import):
                        dependencies.extend({"name": item.name, "line": node.lineno} for item in node.names)
                    elif isinstance(node, ast.ImportFrom):
                        dependencies.append({"name": node.module or "", "line": node.lineno})
                    elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                        item = {"name": node.name, "line": node.lineno}
                        if node.name == "main" or any(
                            isinstance(decorator, ast.Name) and decorator.id == "app"
                            for decorator in node.decorator_list
                        ):
                            entrypoints.append(item)
                        if node.name.startswith("test_"):
                            tests.append(item)
                        for decorator in node.decorator_list:
                            if isinstance(decorator, ast.Call) and isinstance(
                                decorator.func, ast.Attribute
                            ) and decorator.func.attr in {"get", "post", "put", "delete", "route"}:
                                endpoints.append(item | {"method": decorator.func.attr})
        else:
            for match in re.finditer(r"^\s*(?:import|from)\s+([^\s;]+)|require\(['\"]([^'\"]+)", text, re.MULTILINE):
                dependencies.append({"name": match.group(1) or match.group(2), "line": text.count("\n", 0, match.start()) + 1})
            tests = [{"name": item["name"], "line": item["line"]} for item in symbols if str(item["name"]).startswith("test")]
        return {
            "dependencies": dependencies,
            "entrypoints": entrypoints,
            "api_endpoints": endpoints,
            "classes": [item for item in symbols if item["kind"] == "class"],
            "functions": [item for item in symbols if item["kind"] in {"function", "symbol"}],
            "tests": tests,
        }


class CsvParser(SourceParser):
    suffixes = frozenset({".csv"})

    def __init__(self, rows_per_fragment: int = 25, max_rows: int = 200_000, max_columns: int = 500) -> None:
        self.rows_per_fragment = max(1, rows_per_fragment)
        self.max_rows = max_rows
        self.max_columns = max_columns

    def parse(self, source: Source) -> ParseResult:
        path = _path(source)
        raw = path.read_bytes()
        text, encoding = _decode(raw)
        sample = text[:8192]
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        reader = csv.reader(io.StringIO(text, newline=""), dialect)
        result = ParseResult(source.id, metadata={"encoding": encoding, "delimiter": dialect.delimiter})
        batch: list[list[str]] = []
        batch_start = 1
        header: list[str] | None = None
        try:
            for row_index, row in enumerate(reader, start=1):
                if row_index > self.max_rows:
                    result.warnings.append(f"row-limit:{self.max_rows}")
                    break
                if len(row) > self.max_columns:
                    result.warnings.append(f"columns-truncated:row-{row_index}")
                    row = row[: self.max_columns]
                cleaned = [_clean_cell(value) for value in row]
                if header is None:
                    header = cleaned
                batch.append(cleaned)
                if len(batch) >= self.rows_per_fragment:
                    self._append_batch(result, source, path, batch, batch_start, row_index, header)
                    batch = []
                    batch_start = row_index + 1
            if batch:
                end = batch_start + len(batch) - 1
                self._append_batch(result, source, path, batch, batch_start, end, header)
        except csv.Error as error:
            result.warnings.append(f"csv-error:{error}")
        if not result.fragments:
            result.warnings.append("empty-csv")
        return result

    @staticmethod
    def _append_batch(
        result: ParseResult,
        source: Source,
        path: Path,
        rows: list[list[str]],
        start: int,
        end: int,
        header: list[str] | None,
    ) -> None:
        rendered = "\n".join(" | ".join(row) for row in rows)
        result.fragments.append(
            fragment(
                source_id=source.id,
                content=rendered,
                source_locator=locator(
                    source_id=source.id,
                    path=path,
                    row=start,
                    line_start=start,
                    line_end=end,
                    cell_range=f"A{start}:{_column_name(max(map(len, rows), default=1))}{end}",
                    details={"row_end": end},
                ),
                metadata={"header": header or [], "row_count": len(rows)},
                ordinal=f"csv:{start}-{end}",
            )
        )


class XlsxParser(SourceParser):
    suffixes = frozenset({".xlsx"})

    def __init__(self, rows_per_fragment: int = 25, max_rows_per_sheet: int = 200_000) -> None:
        self.rows_per_fragment = max(1, rows_per_fragment)
        self.max_rows_per_sheet = max_rows_per_sheet

    def parse(self, source: Source) -> ParseResult:
        try:
            from openpyxl import load_workbook
        except ImportError:
            return ParseResult(source.id, warnings=[str(OptionalDependencyError("openpyxl", "XLSX"))])

        path = _path(source)
        result = ParseResult(source.id, metadata={"kind": "workbook"})
        try:
            workbook = load_workbook(path, read_only=False, data_only=False, keep_links=False)
            cached_workbook = load_workbook(path, read_only=False, data_only=True, keep_links=False)
        except Exception as error:  # openpyxl exposes several version-specific errors
            result.warnings.append(f"xlsx-error:{type(error).__name__}:{error}")
            return result
        try:
            result.metadata["sheets"] = list(workbook.sheetnames)
            result.metadata["named_ranges"] = sorted(workbook.defined_names)
            for worksheet in workbook.worksheets:
                cached_sheet = cached_workbook[worksheet.title]
                result.metadata.setdefault("merged_cells", {})[worksheet.title] = [
                    str(item) for item in worksheet.merged_cells.ranges
                ]
                batch: list[list[str]] = []
                batch_cells: list[list[dict[str, Any]]] = []
                batch_start = 1
                for row_index, row in enumerate(worksheet.iter_rows(), start=1):
                    if row_index > self.max_rows_per_sheet:
                        result.warnings.append(f"row-limit:{worksheet.title}:{self.max_rows_per_sheet}")
                        break
                    batch.append([_clean_cell(cell.value) for cell in row])
                    batch_cells.append(
                        [
                            {
                                "address": cell.coordinate,
                                "raw_value": _clean_cell(cell.value),
                                "formula": cell.value if cell.data_type == "f" else None,
                                "cached_value": _clean_cell(cached_sheet[cell.coordinate].value),
                                "number_format": cell.number_format,
                                "data_type": cell.data_type,
                            }
                            for cell in row
                        ]
                    )
                    if len(batch) >= self.rows_per_fragment:
                        self._append_batch(
                            result, source, path, worksheet.title, batch, batch_cells, batch_start, row_index
                        )
                        batch = []
                        batch_cells = []
                        batch_start = row_index + 1
                if batch:
                    self._append_batch(
                        result,
                        source,
                        path,
                        worksheet.title,
                        batch,
                        batch_cells,
                        batch_start,
                        batch_start + len(batch) - 1,
                    )
        finally:
            workbook.close()
            cached_workbook.close()
        if not result.fragments:
            result.warnings.append("empty-workbook")
        return result

    @staticmethod
    def _append_batch(
        result: ParseResult,
        source: Source,
        path: Path,
        sheet: str,
        rows: list[list[str]],
        cells: list[list[dict[str, Any]]],
        start: int,
        end: int,
    ) -> None:
        width = max(map(len, rows), default=1)
        content = "\n".join(" | ".join(row) for row in rows)
        result.fragments.append(
            fragment(
                source_id=source.id,
                content=content,
                source_locator=locator(
                    source_id=source.id,
                    path=path,
                    sheet=sheet,
                    row=start,
                    cell_range=f"A{start}:{_column_name(width)}{end}",
                    details={"row_end": end},
                ),
                metadata={"row_count": len(rows), "column_count": width, "cells": cells},
                ordinal=f"xlsx:{sheet}:{start}-{end}",
            )
        )


class DocxParser(SourceParser):
    suffixes = frozenset({".docx"})

    def parse(self, source: Source) -> ParseResult:
        try:
            from docx import Document
        except ImportError:
            return ParseResult(source.id, warnings=[str(OptionalDependencyError("python-docx", "DOCX"))])
        path = _path(source)
        result = ParseResult(source.id, metadata={"kind": "docx"})
        try:
            document = Document(str(path))
        except Exception as error:
            result.warnings.append(f"docx-error:{type(error).__name__}:{error}")
            return result
        properties = document.core_properties
        result.metadata.update(
            {
                "properties": {
                    "title": properties.title or "",
                    "author": properties.author or "",
                    "subject": properties.subject or "",
                    "keywords": properties.keywords or "",
                },
                "sections": [
                    {
                        "top_margin": section.top_margin.pt if section.top_margin else None,
                        "bottom_margin": section.bottom_margin.pt if section.bottom_margin else None,
                        "left_margin": section.left_margin.pt if section.left_margin else None,
                        "right_margin": section.right_margin.pt if section.right_margin else None,
                        "header_paragraphs": len(section.header.paragraphs),
                        "footer_paragraphs": len(section.footer.paragraphs),
                    }
                    for section in document.sections
                ],
                "relationships": [
                    {"id": key, "type": relation.reltype, "target": str(relation.target_ref)}
                    for key, relation in document.part.rels.items()
                ],
            }
        )
        current_section: str | None = None
        for index, paragraph in enumerate(document.paragraphs, start=1):
            content = paragraph.text.strip()
            style = paragraph.style.name if paragraph.style else ""
            if style.casefold().startswith(("heading", "заголов")) and content:
                current_section = content
            if not content:
                continue
            result.fragments.append(
                fragment(
                    source_id=source.id,
                    content=content,
                    source_locator=locator(
                        source_id=source.id,
                        path=path,
                        section=current_section,
                        details={"paragraph": index},
                    ),
                    metadata={"style": style, "kind": "paragraph"},
                    ordinal=f"paragraph:{index}",
                )
            )
        for table_index, table in enumerate(document.tables, start=1):
            rows = [[_clean_cell(cell.text) for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            content = "\n".join(" | ".join(row) for row in rows)
            result.fragments.append(
                fragment(
                    source_id=source.id,
                    content=content,
                    source_locator=locator(
                        source_id=source.id,
                        path=path,
                        section=current_section,
                        cell_range=f"A1:{_column_name(max(map(len, rows), default=1))}{len(rows)}",
                        details={"table": table_index},
                    ),
                    metadata={"kind": "table", "rows": len(rows)},
                    ordinal=f"table:{table_index}",
                )
            )
        result.metadata["inline_shapes"] = len(document.inline_shapes)
        body_xml = document.element.body.xml
        result.metadata["bookmarks"] = body_xml.count("bookmarkStart")
        result.metadata["formulas"] = body_xml.count("oMath")
        result.metadata["hyperlinks"] = body_xml.count("hyperlink")
        if not result.fragments:
            result.warnings.append("empty-docx")
        return result


class PdfParser(SourceParser):
    suffixes = frozenset({".pdf"})

    def __init__(self, max_pages: int = 2_000, vision: VisionOCRPort | None = None) -> None:
        self.max_pages = max_pages
        self.vision = vision

    def parse(self, source: Source) -> ParseResult:
        try:
            from pypdf import PdfReader
        except ImportError:
            return ParseResult(source.id, warnings=[str(OptionalDependencyError("pypdf", "PDF"))])
        path = _path(source)
        result = ParseResult(source.id, metadata={"kind": "pdf"})
        try:
            reader = PdfReader(path, strict=False)
            if reader.is_encrypted:
                try:
                    if reader.decrypt("") == 0:
                        result.warnings.append("encrypted-pdf")
                        return result
                except Exception:
                    result.warnings.append("encrypted-pdf")
                    return result
            result.metadata["page_count"] = len(reader.pages)
            for page_number, page in enumerate(reader.pages[: self.max_pages], start=1):
                try:
                    content = (page.extract_text() or "").strip()
                except Exception as error:
                    result.warnings.append(f"page-extraction-error:{page_number}:{type(error).__name__}")
                    continue
                if not content:
                    result.warnings.append(f"ocr-required:page-{page_number}")
                    self._ocr_page(result, source, path, page_number)
                    continue
                result.fragments.append(
                    fragment(
                        source_id=source.id,
                        content=content,
                        source_locator=locator(source_id=source.id, path=path, page=page_number),
                        metadata={"kind": "page"},
                        ordinal=f"page:{page_number}",
                    )
                )
            if len(reader.pages) > self.max_pages:
                result.warnings.append(f"page-limit:{self.max_pages}")
        except Exception as error:
            result.warnings.append(f"pdf-error:{type(error).__name__}:{error}")
        return result

    def _ocr_page(self, result: ParseResult, source: Source, path: Path, page_number: int) -> None:
        if self.vision is None:
            return
        try:
            fitz = __import__("fitz")
        except ImportError:
            result.warnings.append(f"ocr-render-unavailable:page-{page_number}")
            return
        try:
            with fitz.open(path) as document:
                rendered = document[page_number - 1].get_pixmap(dpi=200, alpha=False).tobytes("png")
            ocr = self.vision.recognize(rendered, page_number=page_number)
        except Exception as error:
            result.warnings.append(f"ocr-error:page-{page_number}:{type(error).__name__}")
            return
        if not ocr.text.strip():
            result.warnings.append(f"ocr-empty:page-{page_number}")
            return
        result.fragments.append(
            fragment(
                source_id=source.id,
                content=ocr.text,
                source_locator=locator(source_id=source.id, path=path, page=page_number),
                metadata={
                    "kind": "ocr-page",
                    "ocr_confidence": ocr.confidence,
                    "tables": list(ocr.tables),
                    "captions": list(ocr.captions),
                    "low_confidence_numbers": list(ocr.low_confidence_numbers),
                },
                ordinal=f"ocr-page:{page_number}",
            )
        )


class ImageParser(SourceParser):
    suffixes = IMAGE_SUFFIXES

    def __init__(self, vision: VisionOCRPort | None = None) -> None:
        self.vision = vision

    def parse(self, source: Source) -> ParseResult:
        try:
            from PIL import Image
        except ImportError:
            return ParseResult(source.id, warnings=[str(OptionalDependencyError("Pillow", "image"))])
        path = _path(source)
        result = ParseResult(source.id, metadata={"kind": "image", "vision_required": True})
        try:
            with Image.open(path) as image:
                image.verify()
            with Image.open(path) as image:
                width, height = image.size
                image_format = image.format
                mode = image.mode
            result.metadata.update({"width": width, "height": height, "format": image_format, "mode": mode})
            description = f"Image {path.name}: {width}x{height}, format {image_format}, mode {mode}"
            result.fragments.append(
                fragment(
                    source_id=source.id,
                    content=description,
                    source_locator=locator(source_id=source.id, path=path),
                    metadata={"kind": "image-metadata", "vision_required": True},
                    ordinal="image",
                )
            )
            if self.vision is not None:
                ocr = self.vision.recognize(path.read_bytes(), page_number=1)
                if ocr.text.strip():
                    result.fragments.append(
                        fragment(
                            source_id=source.id,
                            content=ocr.text,
                            source_locator=locator(source_id=source.id, path=path, page=1),
                            metadata={
                                "kind": "ocr-image",
                                "ocr_confidence": ocr.confidence,
                                "tables": list(ocr.tables),
                                "captions": list(ocr.captions),
                                "low_confidence_numbers": list(ocr.low_confidence_numbers),
                            },
                            ordinal="ocr-image",
                        )
                    )
        except Exception as error:
            result.warnings.append(f"image-error:{type(error).__name__}:{error}")
        return result


class ParserRegistry:
    """Select a parser without importing optional libraries at module import."""

    def __init__(
        self,
        parsers: Iterable[SourceParser] | None = None,
        *,
        vision: VisionOCRPort | None = None,
    ) -> None:
        self.parsers = list(
            parsers
            or (
                DocxParser(),
                PdfParser(vision=vision),
                XlsxParser(),
                CsvParser(),
                ImageParser(vision=vision),
                CodeParser(),
                TextParser(),
            )
        )

    def parser_for(self, source: Source) -> SourceParser:
        for parser in self.parsers:
            if parser.supports(source):
                return parser
        raise UnsupportedSourceError(f"Unsupported input type: {_path(source).suffix or '<none>'}")

    def parse(self, source: Source) -> ParseResult:
        return self.parser_for(source).parse(source)


def parse_source(source: Source, registry: ParserRegistry | None = None) -> ParseResult:
    return (registry or ParserRegistry()).parse(source)


__all__ = [
    "CodeParser",
    "CsvParser",
    "DocxParser",
    "ImageParser",
    "ParserRegistry",
    "PdfParser",
    "SourceParser",
    "TextParser",
    "XlsxParser",
    "parse_source",
]
