"""Render a domain ``Manuscript`` into an editable, field-aware DOCX."""

from __future__ import annotations

import hashlib
import os
import re
import tempfile
import zipfile
from collections.abc import Callable, Mapping
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path, PurePosixPath
from typing import Any
from xml.etree import ElementTree

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from papercraft.domain import (
    AppendixBlock,
    BibliographyEntry,
    ChartBlock,
    Citation,
    CitationBlock,
    CodeListingBlock,
    Dataset,
    DiagramBlock,
    FigureBlock,
    FormulaBlock,
    HeadingBlock,
    Manuscript,
    PageBreakBlock,
    ParagraphBlock,
    TableBlock,
    TemplateApplicationPlan,
)


class DocxRenderError(RuntimeError):
    pass


@dataclass(frozen=True, slots=True)
class TemplateInspection:
    path: Path
    sha256: str
    part_count: int
    uncompressed_bytes: int
    external_hyperlinks: int


def inspect_docx_template(path: str | os.PathLike[str]) -> TemplateInspection:
    """Reject active/remote package features before opening a user template."""

    template = Path(path).expanduser().resolve(strict=True)
    if template.suffix.casefold() != ".docx" or not zipfile.is_zipfile(template):
        raise DocxRenderError("the institution template must be a valid DOCX package")
    forbidden_parts = (
        "vbaproject.bin",
        "word/activex/",
        "word/embeddings/",
        "word/altchunks/",
    )
    total_uncompressed = 0
    external_hyperlinks = 0
    with zipfile.ZipFile(template) as archive:
        entries = archive.infolist()
        if len(entries) > 2_000:
            raise DocxRenderError("template package contains too many parts")
        for entry in entries:
            normalized = entry.filename.replace("\\", "/")
            parts = PurePosixPath(normalized).parts
            if normalized.startswith("/") or ".." in parts:
                raise DocxRenderError("template package contains an unsafe path")
            lowered = normalized.casefold()
            if any(item in lowered for item in forbidden_parts):
                raise DocxRenderError(f"template contains active or embedded content: {normalized}")
            total_uncompressed += entry.file_size
            if (
                entry.file_size > 1_000_000
                and entry.compress_size > 0
                and entry.file_size / entry.compress_size > 500
            ):
                raise DocxRenderError("template package has an unsafe compression ratio")
        if total_uncompressed > 250 * 1024 * 1024:
            raise DocxRenderError("template package is too large after decompression")

        for entry in entries:
            lowered = entry.filename.casefold()
            if lowered.endswith(".rels"):
                try:
                    root = ElementTree.fromstring(archive.read(entry))
                except ElementTree.ParseError as exc:
                    raise DocxRenderError(f"template relationship XML is invalid: {entry.filename}") from exc
                for relationship in root:
                    if relationship.attrib.get("TargetMode", "").casefold() != "external":
                        continue
                    kind = relationship.attrib.get("Type", "").rsplit("/", 1)[-1].casefold()
                    target = relationship.attrib.get("Target", "")
                    if kind != "hyperlink" or not target.casefold().startswith(
                        ("https://", "http://", "mailto:")
                    ):
                        raise DocxRenderError(
                            f"template contains a forbidden external relationship: {kind or 'unknown'}"
                        )
                    external_hyperlinks += 1
            if lowered.startswith("word/") and lowered.endswith(".xml"):
                payload = archive.read(entry)
                if re.search(
                    rb"\b(?:DDEAUTO|DDE|INCLUDETEXT|INCLUDEPICTURE)\b",
                    payload,
                    flags=re.IGNORECASE,
                ):
                    raise DocxRenderError(f"template contains an active field: {entry.filename}")
                if b"<w:altChunk" in payload or b"<o:OLEObject" in payload:
                    raise DocxRenderError(f"template contains active embedded content: {entry.filename}")
    return TemplateInspection(
        path=template,
        sha256=_sha256(template),
        part_count=len(entries),
        uncompressed_bytes=total_uncompressed,
        external_hyperlinks=external_hyperlinks,
    )


@dataclass(frozen=True, slots=True)
class RenderConfig:
    font_name: str = "Times New Roman"
    body_font_size_pt: float = 14
    heading_1_size_pt: float = 16
    heading_2_size_pt: float = 14
    table_font_size_pt: float = 11
    code_font_name: str = "Consolas"
    code_font_size_pt: float = 9
    line_spacing: float = 1.5
    paragraph_indent_cm: float = 1.25
    margin_left_cm: float = 3.0
    margin_right_cm: float = 1.5
    margin_top_cm: float = 2.0
    margin_bottom_cm: float = 2.0
    header_distance_cm: float = 1.25
    footer_distance_cm: float = 1.25
    include_title_page: bool = True
    include_toc: bool = True
    bibliography_heading: str = "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"
    table_label: str = "Таблица"
    figure_label: str = "Рисунок"
    page_number_alignment: str = "center"
    page_number_position: str = "bottom"
    maximum_image_width_cm: float = 16.5
    maximum_image_height_cm: float = 21.0

    def __post_init__(self) -> None:
        positive = {
            "body_font_size_pt": self.body_font_size_pt,
            "line_spacing": self.line_spacing,
            "margin_left_cm": self.margin_left_cm,
            "margin_right_cm": self.margin_right_cm,
            "margin_top_cm": self.margin_top_cm,
            "margin_bottom_cm": self.margin_bottom_cm,
            "maximum_image_width_cm": self.maximum_image_width_cm,
        }
        if any(value <= 0 for value in positive.values()):
            raise ValueError("render dimensions and font sizes must be positive")
        if self.page_number_alignment not in {"left", "center", "right"}:
            raise ValueError("page_number_alignment must be left, center or right")
        if self.page_number_position not in {"top", "bottom"}:
            raise ValueError("page_number_position must be top or bottom")


@dataclass(frozen=True, slots=True)
class TitlePageInfo:
    ministry: str = ""
    university: str = ""
    faculty: str = ""
    department: str = ""
    work_type: str = "АКАДЕМИЧЕСКАЯ РАБОТА"
    subject: str = ""
    topic: str = ""
    student: str = ""
    supervisor: str = ""
    city: str = ""
    year: int = field(default_factory=lambda: date.today().year)

    @classmethod
    def from_mapping(cls, value: Mapping[str, Any], *, fallback_topic: str) -> TitlePageInfo:
        aliases = {
            "student_info": "student",
            "teacher_info": "supervisor",
        }
        accepted = {item.name for item in cls.__dataclass_fields__.values()}
        data: dict[str, Any] = {}
        for key, raw_value in value.items():
            normalized = aliases.get(key, key)
            if normalized in accepted:
                data[normalized] = raw_value
        data.setdefault("topic", fallback_topic)
        if "year" in data:
            try:
                data["year"] = int(data["year"])
            except (TypeError, ValueError):
                data["year"] = date.today().year
        return cls(**data)


@dataclass(frozen=True, slots=True)
class DocxRenderResult:
    path: Path
    sha256: str
    block_count: int
    unresolved_artifact_ids: tuple[str, ...]
    warnings: tuple[str, ...]


ArtifactResolver = Mapping[str, str | os.PathLike[str]] | Callable[[str], str | os.PathLike[str] | None]


class DocxRenderer:
    """Translate manuscript blocks to OpenXML without mixing generation logic."""

    def __init__(self, config: RenderConfig | None = None) -> None:
        self.config = config or RenderConfig()
        self._warnings: list[str] = []
        self._unresolved: list[str] = []
        self._bookmark_counter = 0
        self._sequence_counters: dict[str, int] = {}

    def render(
        self,
        manuscript: Manuscript,
        output_path: str | os.PathLike[str],
        *,
        template_path: str | os.PathLike[str] | None = None,
        template_plan: TemplateApplicationPlan | None = None,
        artifact_paths: ArtifactResolver | None = None,
        datasets: Mapping[str, Dataset] | None = None,
        citations: Mapping[str, Citation] | None = None,
        title_page: TitlePageInfo | Mapping[str, Any] | None = None,
    ) -> DocxRenderResult:
        output = Path(output_path).expanduser().resolve()
        if output.suffix.lower() != ".docx":
            raise DocxRenderError("output path must have a .docx extension")
        self._warnings = []
        self._unresolved = []
        self._bookmark_counter = 0
        self._sequence_counters = {}
        datasets = datasets or {}
        citations = citations or {}

        if template_plan is not None and template_path is None:
            raise DocxRenderError("a TemplateApplicationPlan requires a template_path")
        template_inspection: TemplateInspection | None = None
        if template_path is not None:
            template_inspection = inspect_docx_template(template_path)
            template = template_inspection.path
            if output == template:
                raise DocxRenderError("output must not overwrite the retained institution template")
            document = Document(str(template))
            if template_plan is not None:
                available_styles = {style.name for style in document.styles}
                requested_styles = {
                    *template_plan.use_styles,
                    *template_plan.section_style_map.values(),
                }
                if missing_styles := requested_styles - available_styles:
                    raise DocxRenderError(
                        "template plan references missing styles: " + ", ".join(sorted(missing_styles))
                    )
            self._warnings.append(
                f"Institution template safety-checked and retained: {template.name}"
            )
        else:
            document = Document()
        self._configure_document(document, manuscript, preserve_template=template_path is not None)
        resolved_title = self._title_page(manuscript, title_page)
        if self.config.include_title_page and template_path is None:
            self._render_title_page(document, resolved_title)
        if self.config.include_toc:
            self._render_toc(document)

        for block in manuscript.blocks:
            self._render_block(document, block, artifact_paths, datasets, citations, template_plan)
        if manuscript.bibliography:
            self._render_bibliography(document, manuscript.bibliography)

        output.parent.mkdir(parents=True, exist_ok=True)
        descriptor, temporary_name = tempfile.mkstemp(
            prefix=f".{output.stem}.", suffix=".docx", dir=output.parent
        )
        os.close(descriptor)
        temporary = Path(temporary_name)
        try:
            document.save(str(temporary))
            if temporary.stat().st_size < 1_000:
                raise DocxRenderError("python-docx produced an unexpectedly small document")
            os.replace(temporary, output)
        finally:
            temporary.unlink(missing_ok=True)
        if template_inspection is not None and _sha256(template_inspection.path) != template_inspection.sha256:
            output.unlink(missing_ok=True)
            raise DocxRenderError("the retained institution template changed during rendering")
        return DocxRenderResult(
            path=output,
            sha256=_sha256(output),
            block_count=len(manuscript.blocks),
            unresolved_artifact_ids=tuple(dict.fromkeys(self._unresolved)),
            warnings=tuple(self._warnings),
        )

    def _configure_document(
        self,
        document: Any,
        manuscript: Manuscript,
        *,
        preserve_template: bool,
    ) -> None:
        document.core_properties.title = manuscript.title
        document.core_properties.subject = "PaperCraft AI Studio generated manuscript"
        document.core_properties.keywords = "academic, evidence-backed"
        document.core_properties.comments = f"Project {manuscript.project_id}; revision {manuscript.revision}"
        if not preserve_template:
            for section in document.sections:
                section.left_margin = Cm(self.config.margin_left_cm)
                section.right_margin = Cm(self.config.margin_right_cm)
                section.top_margin = Cm(self.config.margin_top_cm)
                section.bottom_margin = Cm(self.config.margin_bottom_cm)
                section.header_distance = Cm(self.config.header_distance_cm)
                section.footer_distance = Cm(self.config.footer_distance_cm)
                section.different_first_page_header_footer = True
                page_container = (
                    section.header if self.config.page_number_position == "top" else section.footer
                )
                self._add_page_field(page_container.paragraphs[0])

            normal = document.styles["Normal"]
            _set_style_font(normal, self.config.font_name, self.config.body_font_size_pt)
            normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            normal.paragraph_format.line_spacing = self.config.line_spacing
            normal.paragraph_format.first_line_indent = Cm(self.config.paragraph_indent_cm)
            normal.paragraph_format.space_before = Pt(0)
            normal.paragraph_format.space_after = Pt(0)

            for level in range(1, 7):
                name = f"Heading {level}"
                if name not in document.styles:
                    continue
                style = document.styles[name]
                size = self.config.heading_1_size_pt if level == 1 else self.config.heading_2_size_pt
                _set_style_font(style, self.config.font_name, size, bold=True)
                style.font.color.rgb = RGBColor(0, 0, 0)
                style.paragraph_format.first_line_indent = Cm(0)
                style.paragraph_format.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
                )
                style.paragraph_format.space_before = Pt(12)
                style.paragraph_format.space_after = Pt(6)
                style.paragraph_format.keep_with_next = True

            if "Caption" in document.styles:
                caption = document.styles["Caption"]
                _set_style_font(caption, self.config.font_name, 12)
                caption.font.italic = False
                caption.font.color.rgb = RGBColor(0, 0, 0)

        settings = document.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    def _title_page(
        self,
        manuscript: Manuscript,
        supplied: TitlePageInfo | Mapping[str, Any] | None,
    ) -> TitlePageInfo:
        if isinstance(supplied, TitlePageInfo):
            return supplied
        if supplied is not None:
            return TitlePageInfo.from_mapping(supplied, fallback_topic=manuscript.title)
        metadata_value = manuscript.metadata.get("title_page")
        if isinstance(metadata_value, dict):
            return TitlePageInfo.from_mapping(metadata_value, fallback_topic=manuscript.title)
        return TitlePageInfo(topic=manuscript.title)

    def _render_title_page(self, document: Any, title: TitlePageInfo) -> None:
        top_lines = [title.ministry, title.university, title.faculty, title.department]
        paragraph = document.add_paragraph()
        _plain_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
        paragraph.paragraph_format.space_after = Pt(50)
        run = paragraph.add_run("\n".join(line for line in top_lines if line))
        _set_run_font(run, self.config.font_name, 12, bold=True)

        title_paragraph = document.add_paragraph()
        _plain_paragraph(title_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
        title_paragraph.paragraph_format.space_before = Pt(30)
        title_paragraph.paragraph_format.space_after = Pt(65)
        work_run = title_paragraph.add_run(title.work_type.upper())
        _set_run_font(work_run, self.config.font_name, 16, bold=True)
        if title.subject:
            subject_run = title_paragraph.add_run(f"\n{title.subject}")
            _set_run_font(subject_run, self.config.font_name, 14)
        topic_run = title_paragraph.add_run(f"\n\nна тему: «{title.topic}»")
        _set_run_font(topic_run, self.config.font_name, 14, bold=True)

        author_table = document.add_table(rows=1, cols=2)
        author_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        _remove_table_borders(author_table)
        right = author_table.cell(0, 1)
        lines = [title.student, title.supervisor]
        author = right.paragraphs[0]
        _plain_paragraph(author, WD_ALIGN_PARAGRAPH.LEFT)
        author.paragraph_format.line_spacing = 1.15
        author_run = author.add_run("\n\n".join(line for line in lines if line))
        _set_run_font(author_run, self.config.font_name, 12)

        bottom = document.add_paragraph()
        _plain_paragraph(bottom, WD_ALIGN_PARAGRAPH.CENTER)
        bottom.paragraph_format.space_before = Pt(110)
        location = " — ".join(part for part in [title.city, str(title.year)] if part)
        bottom_run = bottom.add_run(location)
        _set_run_font(bottom_run, self.config.font_name, 12)
        bottom.add_run().add_break(WD_BREAK.PAGE)

    def _render_toc(self, document: Any) -> None:
        # Do not use Heading 1 here: otherwise the TOC includes itself.
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.first_line_indent = Cm(0)
        heading_run = heading.add_run("СОДЕРЖАНИЕ")
        _set_run_font(
            heading_run, self.config.font_name, self.config.heading_1_size_pt, bold=True
        )
        paragraph = document.add_paragraph()
        _plain_paragraph(paragraph, WD_ALIGN_PARAGRAPH.LEFT)
        _append_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u', "Обновите оглавление в Word")
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    def _render_block(
        self,
        document: Any,
        block: Any,
        artifact_paths: ArtifactResolver | None,
        datasets: Mapping[str, Dataset],
        citations: Mapping[str, Citation],
        template_plan: TemplateApplicationPlan | None = None,
    ) -> None:
        if isinstance(block, ParagraphBlock):
            style = block.style or (template_plan.section_style_map.get("paragraph") if template_plan else None) or "Normal"
            paragraph = document.add_paragraph(style=style)
            paragraph.add_run(block.text)
            for citation_id in block.citation_ids:
                citation = citations.get(citation_id)
                marker = citation.marker if citation else f"[{citation_id}]"
                paragraph.add_run(f" {marker}")
                if citation is None:
                    self._warnings.append(f"Citation {citation_id} could not be resolved")
            return
        if isinstance(block, HeadingBlock):
            style = (
                template_plan.section_style_map.get(block.section_id or "")
                if template_plan
                else None
            ) or f"Heading {min(block.level, 6)}"
            document.add_paragraph(block.text, style=style)
            return
        if isinstance(block, TableBlock):
            self._render_table(document, block, datasets)
            return
        if isinstance(block, (ChartBlock, DiagramBlock)):
            caption = block.spec.title
            self._render_artifact_image(
                document, block.artifact_id, caption, artifact_paths, self.config.figure_label
            )
            return
        if isinstance(block, FigureBlock):
            self._render_artifact_image(
                document, block.artifact_id, block.caption, artifact_paths, self.config.figure_label
            )
            return
        if isinstance(block, FormulaBlock):
            paragraph = document.add_paragraph()
            _plain_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
            _append_omml_formula(paragraph, block.spec.expression)
            if block.spec.label:
                paragraph.add_run(f"    ({block.spec.label})")
            return
        if isinstance(block, CodeListingBlock):
            self._render_code(document, block)
            return
        if isinstance(block, CitationBlock):
            paragraph = document.add_paragraph(style="Normal")
            citation = citations.get(block.citation_id)
            marker = citation.marker if citation else f"[{block.citation_id}]"
            paragraph.add_run(f"{block.text} {marker}".strip())
            if citation is None:
                self._warnings.append(f"Citation {block.citation_id} could not be resolved")
            return
        if isinstance(block, PageBreakBlock):
            document.add_page_break()
            return
        if isinstance(block, AppendixBlock):
            document.add_page_break()
            document.add_paragraph(block.title, style="Heading 1")
            for nested in block.blocks:
                self._render_block(document, nested, artifact_paths, datasets, citations, template_plan)
            return
        self._warnings.append(f"Unsupported manuscript block: {type(block).__name__}")

    def _render_table(
        self, document: Any, block: TableBlock, datasets: Mapping[str, Dataset]
    ) -> None:
        spec = block.spec
        headers = list(spec.headers)
        rows = [list(row) for row in spec.rows]
        if spec.dataset_id and (not headers or not rows):
            dataset = datasets.get(spec.dataset_id)
            if dataset is None:
                self._warnings.append(f"Dataset {spec.dataset_id} could not be resolved")
            else:
                headers = [column.name for column in dataset.columns]
                rows = [[row.get(header, "") for header in headers] for row in dataset.rows]
        if not headers:
            headers = ["Данные"]
        if any(len(row) != len(headers) for row in rows):
            raise DocxRenderError(f"table {block.id} rows do not match headers")

        use_landscape = len(headers) > 6 or (
            bool(spec.column_widths) and sum(float(value) for value in spec.column_widths) > 17
        )
        previous_section = document.sections[-1]
        previous_orientation = previous_section.orientation
        previous_page_width = previous_section.page_width
        previous_page_height = previous_section.page_height
        previous_left_margin = previous_section.left_margin
        previous_right_margin = previous_section.right_margin
        previous_top_margin = previous_section.top_margin
        previous_bottom_margin = previous_section.bottom_margin
        previous_header_distance = previous_section.header_distance
        previous_footer_distance = previous_section.footer_distance
        if use_landscape:
            landscape = document.add_section(WD_SECTION_START.NEW_PAGE)
            landscape.orientation = WD_ORIENT.LANDSCAPE
            landscape.page_width = previous_page_height
            landscape.page_height = previous_page_width
            landscape.left_margin = previous_left_margin
            landscape.right_margin = previous_right_margin
            landscape.top_margin = previous_top_margin
            landscape.bottom_margin = previous_bottom_margin
            landscape.header_distance = previous_header_distance
            landscape.footer_distance = previous_footer_distance
            landscape.different_first_page_header_footer = False
            landscape.header.is_linked_to_previous = True
            landscape.footer.is_linked_to_previous = True
            self._warnings.append(f"Wide table {block.id} placed in a landscape section")
        self._add_caption(document, self.config.table_label, "Table", spec.caption, above=True)
        table = document.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        table.autofit = not bool(spec.column_widths)
        if spec.column_widths and len(spec.column_widths) != len(headers):
            self._warnings.append(f"Table {block.id} column widths were ignored")
        for column_index, header in enumerate(headers):
            cell = table.rows[0].cells[column_index]
            cell.text = str(header)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _shade_cell(cell, "E7EAF0")
            _format_cell(cell, self.config, bold=True, centered=True)
            if len(spec.column_widths) == len(headers):
                cell.width = Cm(float(spec.column_widths[column_index]))
        _repeat_table_header(table.rows[0])
        _prevent_row_split(table.rows[0])
        for row_index, values in enumerate(rows, start=1):
            _prevent_row_split(table.rows[row_index])
            for column_index, value in enumerate(values):
                cell = table.rows[row_index].cells[column_index]
                cell.text = "" if value is None else str(value)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _format_cell(cell, self.config, centered=column_index > 0)
                if len(spec.column_widths) == len(headers):
                    cell.width = Cm(float(spec.column_widths[column_index]))
        document.add_paragraph().paragraph_format.space_after = Pt(4)
        if use_landscape:
            restored = document.add_section(WD_SECTION_START.NEW_PAGE)
            restored.orientation = previous_orientation
            restored.page_width = previous_page_width
            restored.page_height = previous_page_height
            restored.left_margin = previous_left_margin
            restored.right_margin = previous_right_margin
            restored.top_margin = previous_top_margin
            restored.bottom_margin = previous_bottom_margin
            restored.header_distance = previous_header_distance
            restored.footer_distance = previous_footer_distance
            restored.different_first_page_header_footer = False
            restored.header.is_linked_to_previous = True
            restored.footer.is_linked_to_previous = True

    def _render_artifact_image(
        self,
        document: Any,
        artifact_id: str | None,
        caption: str,
        artifact_paths: ArtifactResolver | None,
        label: str,
    ) -> None:
        path = self._resolve_artifact(artifact_id, artifact_paths)
        if path is None or not path.is_file():
            missing_id = artifact_id or "generated-image"
            self._unresolved.append(missing_id)
            placeholder = document.add_paragraph(f"[[MISSING ARTIFACT: {missing_id}]]")
            placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            placeholder.runs[0].font.color.rgb = RGBColor(192, 0, 0)
            if caption:
                self._add_caption(document, label, "Figure", caption, above=False)
            return
        image_paragraph = document.add_paragraph()
        _plain_paragraph(image_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
        image_paragraph.paragraph_format.keep_with_next = True
        width_cm, height_cm = _fit_image(
            path, self.config.maximum_image_width_cm, self.config.maximum_image_height_cm
        )
        image_paragraph.add_run().add_picture(
            str(path), width=Cm(width_cm), height=Cm(height_cm)
        )
        self._add_caption(document, label, "Figure", caption, above=False)

    @staticmethod
    def _resolve_artifact(
        artifact_id: str | None, artifact_paths: ArtifactResolver | None
    ) -> Path | None:
        if artifact_id is None or artifact_paths is None:
            return None
        value = artifact_paths(artifact_id) if callable(artifact_paths) else artifact_paths.get(artifact_id)
        return Path(value).expanduser().resolve() if value else None

    def _render_code(self, document: Any, block: CodeListingBlock) -> None:
        if block.caption:
            paragraph = document.add_paragraph()
            _plain_paragraph(paragraph, WD_ALIGN_PARAGRAPH.LEFT)
            run = paragraph.add_run(block.caption)
            _set_run_font(run, self.config.font_name, 12)
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        _shade_cell(cell, "F4F4F4")
        paragraph = cell.paragraphs[0]
        _plain_paragraph(paragraph, WD_ALIGN_PARAGRAPH.LEFT)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(block.code)
        _set_run_font(run, self.config.code_font_name, self.config.code_font_size_pt)

    def _render_bibliography(
        self, document: Any, entries: list[BibliographyEntry]
    ) -> None:
        document.add_page_break()
        document.add_paragraph(self.config.bibliography_heading, style="Heading 1")
        for index, entry in enumerate(entries, start=1):
            paragraph = document.add_paragraph(style="Normal")
            paragraph.paragraph_format.left_indent = Cm(0.75)
            paragraph.paragraph_format.first_line_indent = Cm(-0.75)
            paragraph.add_run(f"{index}. {_bibliography_text(entry)}")

    def _add_caption(
        self,
        document: Any,
        label: str,
        sequence_name: str,
        caption: str,
        *,
        above: bool,
    ) -> None:
        paragraph = document.add_paragraph(style="Caption")
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if above else WD_ALIGN_PARAGRAPH.CENTER
        )
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.keep_with_next = above
        paragraph.add_run(f"{label} ")
        self._bookmark_counter += 1
        sequence_number = self._sequence_counters.get(sequence_name, 0) + 1
        self._sequence_counters[sequence_name] = sequence_number
        bookmark_id = str(self._bookmark_counter)
        bookmark_name = f"pc_{sequence_name}_{sequence_number}"
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bookmark_id)
        start.set(qn("w:name"), bookmark_name)
        paragraph._p.append(start)
        _append_field(paragraph, f"SEQ {sequence_name} \\* ARABIC", str(sequence_number))
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bookmark_id)
        paragraph._p.append(end)
        if caption:
            paragraph.add_run(f" – {caption}")

    def _add_page_field(self, paragraph: Any) -> None:
        alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[self.config.page_number_alignment]
        _plain_paragraph(paragraph, alignment)
        _append_field(paragraph, "PAGE", "1")


def _set_style_font(style: Any, name: str, size: float, *, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.element.rPr.rFonts.set(qn("w:cs"), name)


def _set_run_font(run: Any, name: str, size: float, *, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)


def _plain_paragraph(paragraph: Any, alignment: Any) -> None:
    paragraph.alignment = alignment
    paragraph.paragraph_format.first_line_indent = Cm(0)


def _append_field(paragraph: Any, instruction: str, fallback: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, code, separate, text, end):
        run._r.append(element)


def _remove_table_borders(table: Any) -> None:
    properties = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    properties.append(borders)


def _shade_cell(cell: Any, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _format_cell(
    cell: Any, config: RenderConfig, *, bold: bool = False, centered: bool = False
) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
        )
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            _set_run_font(run, config.font_name, config.table_font_size_pt, bold=bold)


def _repeat_table_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _prevent_row_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    properties.append(cannot_split)


def _append_omml_formula(paragraph: Any, expression: str) -> None:
    math_paragraph = OxmlElement("m:oMathPara")
    math = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    text = OxmlElement("m:t")
    text.text = expression
    run.append(text)
    math.append(run)
    math_paragraph.append(math)
    paragraph._p.append(math_paragraph)


def _fit_image(path: Path, maximum_width_cm: float, maximum_height_cm: float) -> tuple[float, float]:
    try:
        from PIL import Image

        with Image.open(path) as image:
            width, height = image.size
    except Exception as exc:
        raise DocxRenderError(f"cannot read image {path}: {exc}") from exc
    if width <= 0 or height <= 0:
        raise DocxRenderError(f"image {path} has invalid dimensions")
    ratio = width / height
    target_width = maximum_width_cm
    target_height = target_width / ratio
    if target_height > maximum_height_cm:
        target_height = maximum_height_cm
        target_width = target_height * ratio
    return target_width, target_height


def _bibliography_text(entry: BibliographyEntry) -> str:
    if entry.citation_text.strip():
        return entry.citation_text.strip()
    authors = ", ".join(entry.authors)
    parts = [f"{authors}." if authors else "", entry.title]
    if entry.publisher:
        parts.append(f"— {entry.publisher}")
    if entry.year:
        parts.append(str(entry.year))
    text = " ".join(part for part in parts if part).rstrip(".") + "."
    if entry.doi:
        text += f" DOI: {entry.doi}."
    if entry.url:
        text += f" URL: {entry.url}"
        if entry.accessed_on:
            text += f" (дата обращения: {entry.accessed_on.strftime('%d.%m.%Y')})"
        text += "."
    return text


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()
