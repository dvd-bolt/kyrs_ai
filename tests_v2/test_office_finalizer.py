from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

import pytest
from docx import Document

from papercraft.infrastructure.render import (
    DocumentFinalizer,
    FinalizationError,
    FinalizationUnavailableError,
)

_PDF = b"%PDF-1.7\nlocal fixture\n"


def _docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Office finalizer fixture")
    document.save(path)


def test_beta_default_uses_libreoffice(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    docx = tmp_path / "source.docx"
    pdf = tmp_path / "source.pdf"
    _docx(docx)
    finalizer = DocumentFinalizer()
    attempts: list[str] = []

    def libreoffice_succeeds(_docx: Path, output: Path | None) -> bool:
        attempts.append("libreoffice")
        assert output is not None
        output.write_bytes(_PDF)
        return True

    def word_must_not_run(_docx: Path, _output: Path | None) -> None:
        attempts.append("word")
        raise AssertionError("the beta default must not invoke Word")

    monkeypatch.setattr(finalizer, "_convert_with_libreoffice", libreoffice_succeeds)
    monkeypatch.setattr(finalizer, "_finalize_with_word", word_must_not_run)

    result = finalizer.finalize(docx, pdf_path=pdf)

    assert attempts == ["libreoffice"]
    assert result.engine == "libreoffice"
    assert result.fields_updated
    assert result.pdf is not None and result.pdf.valid_header


def test_beta_default_never_falls_back_to_word(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    docx = tmp_path / "source.docx"
    pdf = tmp_path / "source.pdf"
    _docx(docx)
    finalizer = DocumentFinalizer()
    attempts: list[str] = []

    def libreoffice_fails(_docx: Path, _output: Path | None) -> bool:
        attempts.append("libreoffice")
        raise FinalizationError("simulated LibreOffice failure")

    def word_must_not_run(_docx: Path, _output: Path | None) -> None:
        attempts.append("word")
        raise AssertionError("the beta default must not fall back to Word")

    monkeypatch.setattr(finalizer, "_convert_with_libreoffice", libreoffice_fails)
    monkeypatch.setattr(finalizer, "_finalize_with_word", word_must_not_run)

    with pytest.raises(FinalizationUnavailableError, match="LibreOffice"):
        finalizer.finalize(docx, pdf_path=pdf)

    assert attempts == ["libreoffice"]


def test_auto_word_fallback_marks_export_outside_libreoffice_beta_gate(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    docx = tmp_path / "source.docx"
    pdf = tmp_path / "source.pdf"
    _docx(docx)
    finalizer = DocumentFinalizer()
    private_value = "C:/Users/private-user/project/secret-source.docx"

    def libreoffice_fails(_docx: Path, _output: Path | None) -> bool:
        raise FinalizationError(f"conversion failed for {private_value}")

    def word_succeeds(_docx: Path, output: Path | None) -> None:
        assert output is not None
        output.write_bytes(_PDF)

    monkeypatch.setattr(finalizer, "_convert_with_libreoffice", libreoffice_fails)
    monkeypatch.setattr(finalizer, "_finalize_with_word", word_succeeds)

    result = finalizer.finalize(docx, pdf_path=pdf, preferred="auto")

    assert result.engine == "word"
    assert any("compatibility fallback" in warning for warning in result.warnings)
    assert all(private_value not in warning for warning in result.warnings)


def test_explicit_libreoffice_missing_executable_is_release_blocking(tmp_path: Path) -> None:
    docx = tmp_path / "source.docx"
    _docx(docx)
    finalizer = DocumentFinalizer(libreoffice_path=tmp_path / "missing-soffice.com")

    with pytest.raises(FinalizationUnavailableError, match="LibreOffice finalization unavailable"):
        finalizer.finalize(docx, preferred="libreoffice")


def test_command_line_fallback_exports_pdf_with_an_isolated_profile(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    docx = tmp_path / "source.docx"
    pdf = tmp_path / "output" / "source.pdf"
    _docx(docx)
    executable = tmp_path / "fake-office" / "soffice.com"
    executable.parent.mkdir()
    executable.write_bytes(b"fixture")
    finalizer = DocumentFinalizer(libreoffice_path=executable)
    observed: dict[str, object] = {}

    def fake_run(arguments: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
        observed["arguments"] = arguments
        observed["environment"] = kwargs["env"]
        output_directory = Path(arguments[arguments.index("--outdir") + 1])
        (output_directory / "source.pdf").write_bytes(_PDF)
        return subprocess.CompletedProcess(arguments, 0, stdout="noisy", stderr="private stderr")

    monkeypatch.setattr(
        "papercraft.infrastructure.render.finalizer.subprocess.run",
        fake_run,
    )

    result = finalizer.finalize(docx, pdf_path=pdf, preferred="libreoffice")

    environment = observed["environment"]
    assert isinstance(environment, dict)
    assert result.engine == "libreoffice"
    assert not result.fields_updated
    assert pdf.read_bytes() == _PDF
    assert "--norestore" in observed["arguments"]
    assert Path(environment["HOME"]).parent.parent.name.startswith("papercraft-lo-")
    assert not Path(environment["HOME"]).exists()
    assert not list(pdf.parent.glob(".*.tmp"))


def test_uno_export_atomically_replaces_the_updated_docx(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    docx = tmp_path / "source.docx"
    pdf = tmp_path / "source.pdf"
    _docx(docx)
    executable = tmp_path / "fake-office" / "soffice.com"
    executable.parent.mkdir()
    executable.write_bytes(b"fixture")
    (executable.parent / "python.exe").write_bytes(b"fixture")
    finalizer = DocumentFinalizer(libreoffice_path=executable)

    def uno_succeeds(
        _executable: Path,
        _python: Path,
        _helper: Path,
        _profile: Path,
        _source: Path,
        generated: Path,
        updated_docx: Path,
        _creationflags: int,
        _environment: dict[str, str],
    ) -> None:
        generated.write_bytes(_PDF)
        updated = Document()
        updated.add_paragraph("Updated by fake UNO")
        updated.save(updated_docx)

    monkeypatch.setattr(finalizer, "_export_with_uno", uno_succeeds)

    result = finalizer.finalize(docx, pdf_path=pdf, preferred="libreoffice")

    assert result.fields_updated
    assert pdf.read_bytes() == _PDF
    assert [paragraph.text for paragraph in Document(docx).paragraphs] == ["Updated by fake UNO"]
    assert not list(tmp_path.glob(".*.tmp"))


def test_profile_cleanup_retries_transient_windows_file_locks(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    profile = tmp_path / "papercraft-lo-profile"
    profile.mkdir()
    (profile / "registry.dat").write_text("fixture", encoding="utf-8")
    original_rmtree = shutil.rmtree
    attempts = 0

    def delayed_rmtree(path: Path) -> None:
        nonlocal attempts
        attempts += 1
        if attempts < 3:
            raise PermissionError("simulated transient LibreOffice handle")
        original_rmtree(path)

    monkeypatch.setattr(
        "papercraft.infrastructure.render.finalizer.shutil.rmtree",
        delayed_rmtree,
    )
    monkeypatch.setattr("papercraft.infrastructure.render.finalizer.time.sleep", lambda _seconds: None)

    DocumentFinalizer._cleanup_working_directory(profile)

    assert attempts == 3
    assert not profile.exists()


def test_failed_cli_conversion_hides_raw_process_diagnostics(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    docx = tmp_path / "source.docx"
    _docx(docx)
    executable = tmp_path / "fake-office" / "soffice.com"
    executable.parent.mkdir()
    executable.write_bytes(b"fixture")
    finalizer = DocumentFinalizer(libreoffice_path=executable)
    secret = "PRIVATE_DOCUMENT_TEXT_AND_PATH"

    def fake_run(arguments: list[str], **_kwargs: object) -> subprocess.CompletedProcess[str]:
        return subprocess.CompletedProcess(arguments, 19, stdout=secret, stderr=secret)

    monkeypatch.setattr(
        "papercraft.infrastructure.render.finalizer.subprocess.run",
        fake_run,
    )

    with pytest.raises(FinalizationUnavailableError) as failure:
        finalizer.finalize(docx, preferred="libreoffice")

    assert secret not in str(failure.value)
    assert "LibreOffice finalization failed" in str(failure.value)


def test_uno_timeout_terminates_listener_and_hides_process_output(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    docx = tmp_path / "source.docx"
    _docx(docx)
    executable = tmp_path / "fake-office" / "soffice.com"
    executable.parent.mkdir()
    executable.write_bytes(b"fixture")
    (executable.parent / "python.exe").write_bytes(b"fixture")
    finalizer = DocumentFinalizer(libreoffice_path=executable, timeout_seconds=1)
    secret = "PRIVATE_HELPER_OUTPUT"

    class Listener:
        terminated = False
        killed = False

        def poll(self) -> None:
            return None

        def terminate(self) -> None:
            self.terminated = True

        def wait(self, timeout: float) -> int:
            if self.terminated or self.killed:
                return 0
            raise subprocess.TimeoutExpired("listener", timeout)

        def kill(self) -> None:
            self.killed = True

    listener = Listener()

    def fake_popen(*_args: object, **_kwargs: object) -> Listener:
        return listener

    def timeout(*_args: object, **_kwargs: object) -> subprocess.CompletedProcess[str]:
        raise subprocess.TimeoutExpired("helper", 1, output=secret, stderr=secret)

    monkeypatch.setattr(
        "papercraft.infrastructure.render.finalizer.subprocess.Popen",
        fake_popen,
    )
    monkeypatch.setattr(
        "papercraft.infrastructure.render.finalizer.subprocess.run",
        timeout,
    )

    with pytest.raises(FinalizationUnavailableError) as failure:
        finalizer.finalize(docx, preferred="libreoffice")

    assert listener.terminated
    assert secret not in str(failure.value)
    assert "timed out" in str(failure.value)
