from __future__ import annotations

import hashlib
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm

from papercraft.domain import HeadingBlock, Manuscript, ParagraphBlock, TemplateApplicationPlan
from papercraft.infrastructure.render import DocxRenderer, DocxRenderError, RenderConfig


def _template(path: Path) -> str:
    document = Document()
    section = document.sections[0]
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.25)
    section.header.paragraphs[0].text = "Университетский шаблон"
    document.add_paragraph("СОХРАНЯЕМАЯ ТИТУЛЬНАЯ СТРАНИЦА")
    document.save(path)
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _manuscript() -> Manuscript:
    return Manuscript(
        project_id="template-safety",
        title="Новая работа",
        blocks=[HeadingBlock(text="ВВЕДЕНИЕ", level=1), ParagraphBlock(text="Проверенный текст.")],
    )


def test_template_is_not_overwritten_and_geometry_is_preserved(tmp_path: Path) -> None:
    template = tmp_path / "institution.docx"
    original_sha = _template(template)
    output = tmp_path / "result.docx"

    result = DocxRenderer(
        RenderConfig(
            include_title_page=False,
            include_toc=False,
            margin_left_cm=1.0,
            margin_right_cm=1.0,
        )
    ).render(_manuscript(), output, template_path=template)

    assert hashlib.sha256(template.read_bytes()).hexdigest() == original_sha
    loaded = Document(output)
    assert loaded.sections[0].left_margin is not None
    assert loaded.sections[0].right_margin is not None
    assert abs(loaded.sections[0].left_margin.cm - 4.0) < 0.02
    assert abs(loaded.sections[0].right_margin.cm - 2.25) < 0.02
    assert loaded.sections[0].header.paragraphs[0].text == "Университетский шаблон"
    text = "\n".join(paragraph.text for paragraph in loaded.paragraphs)
    assert "СОХРАНЯЕМАЯ ТИТУЛЬНАЯ СТРАНИЦА" in text
    assert "Проверенный текст." in text
    assert any("safety-checked" in warning for warning in result.warnings)


def test_template_cannot_be_its_own_output(tmp_path: Path) -> None:
    template = tmp_path / "institution.docx"
    _template(template)
    with pytest.raises(DocxRenderError, match="must not overwrite"):
        DocxRenderer(RenderConfig(include_toc=False)).render(
            _manuscript(),
            template,
            template_path=template,
        )


def test_template_rejects_embedded_active_content(tmp_path: Path) -> None:
    template = tmp_path / "malicious.docx"
    _template(template)
    with zipfile.ZipFile(template, "a") as archive:
        archive.writestr("word/embeddings/payload.bin", b"untrusted")
    with pytest.raises(DocxRenderError, match="active or embedded"):
        DocxRenderer(RenderConfig(include_toc=False)).render(
            _manuscript(),
            tmp_path / "result.docx",
            template_path=template,
        )


def test_template_plan_must_reference_real_plain_styles(tmp_path: Path) -> None:
    template = tmp_path / "institution.docx"
    _template(template)
    with pytest.raises(DocxRenderError, match="missing styles"):
        DocxRenderer(RenderConfig(include_toc=False)).render(
            _manuscript(),
            tmp_path / "result.docx",
            template_path=template,
            template_plan=TemplateApplicationPlan(use_styles=["Not A Real Style"]),
        )
    with pytest.raises(ValueError, match="plain identifiers"):
        TemplateApplicationPlan(use_styles=["<w:style/>"])
