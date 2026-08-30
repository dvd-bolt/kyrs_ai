from __future__ import annotations

import json
import os
import zipfile
from datetime import date
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont, ImageStat

from papercraft.domain import (
    AppendixBlock,
    BibliographyEntry,
    FigureBlock,
    FormulaBlock,
    FormulaSpec,
    HeadingBlock,
    Manuscript,
    ParagraphBlock,
    TableBlock,
    TableSpec,
)
from papercraft.infrastructure.render import DocumentFinalizer, DocxRenderer, RenderConfig

pytestmark = pytest.mark.office


def test_mixed_page_layout_is_structurally_valid_without_live_office(tmp_path: Path) -> None:
    manuscript = Manuscript(
        project_id="office-layout-structure",
        title="LibreOffice layout structure",
        blocks=[
            HeadingBlock(text="INTRODUCTION", level=1),
            ParagraphBlock(text="Portrait content before a wide table."),
            TableBlock(
                spec=TableSpec(
                    caption="Wide table",
                    headers=[f"Column {index}" for index in range(1, 8)],
                    rows=[[f"R{row}C{column}" for column in range(1, 8)] for row in range(1, 4)],
                )
            ),
            HeadingBlock(text="CONCLUSION", level=1),
            ParagraphBlock(text="Portrait content after the wide table."),
        ],
    )
    docx = tmp_path / "layout-structure.docx"
    DocxRenderer(RenderConfig()).render(manuscript, docx)

    loaded = Document(docx)
    assert len(loaded.sections) >= 3
    assert loaded.sections[0].page_height > loaded.sections[0].page_width
    landscape_sections = [section for section in loaded.sections if section.orientation == 1]
    assert landscape_sections
    assert all(section.page_width > section.page_height for section in landscape_sections)
    assert loaded.sections[-1].page_height > loaded.sections[-1].page_width
    assert all(section.header.is_linked_to_previous for section in loaded.sections[1:])
    assert all(section.footer.is_linked_to_previous for section in loaded.sections[1:])

    with zipfile.ZipFile(docx) as archive:
        document_xml = archive.read("word/document.xml")
        footer_parts = [
            archive.read(name)
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        ]
    assert b"tblHeader" in document_xml
    assert b"cantSplit" in document_xml
    assert any(b"PAGE" in item for item in footer_parts)


@pytest.mark.skipif(
    os.getenv("PAPERCRAFT_RUN_OFFICE_TESTS") != "1",
    reason="set PAPERCRAFT_RUN_OFFICE_TESTS=1 to exercise installed LibreOffice",
)
def test_installed_office_updates_docx_and_exports_valid_pdf(tmp_path: Path) -> None:
    manuscript = Manuscript(
        project_id="office-smoke",
        title="PaperCraft Office smoke test",
        blocks=[
            HeadingBlock(text="INTRODUCTION", level=1),
            ParagraphBlock(
                text=(
                    "This document is generated locally to verify Word fields, "
                    "DOCX persistence and PDF export."
                )
            ),
        ],
    )
    docx = tmp_path / "office-smoke.docx"
    pdf = tmp_path / "office-smoke.pdf"
    DocxRenderer(RenderConfig()).render(manuscript, docx)

    finalizer = DocumentFinalizer(timeout_seconds=120)
    if not finalizer.libreoffice_available():
        pytest.skip("LibreOffice is not available")
    result = finalizer.finalize(docx, pdf_path=pdf, preferred="libreoffice", require_pdf=True)

    assert result.engine == "libreoffice"
    assert result.fields_updated
    assert result.pdf is not None
    assert result.pdf.valid_header
    assert result.pdf.size_bytes > 1_000
    assert docx.is_file() and docx.stat().st_size > 1_000


def _append_ref_field(path: Path, bookmark: str) -> None:
    document = Document(path)
    paragraph = document.add_paragraph("Контрольная перекрёстная ссылка: ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = f" REF {bookmark} \\h "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, code, separate, fallback, end):
        run._r.append(element)
    document.save(path)


def _assert_rendered_pdf_feature_matrix(pdf_path: Path, preview_root: Path) -> dict[str, object]:
    """Check the actual LibreOffice PDF, then retain page PNGs for visual review.

    OOXML checks above prove that the renderer requested the right fields and
    sections.  This check deliberately inspects the *post-LibreOffice* PDF so
    a regression in field updates, Cyrillic glyphs, table layout, or section
    orientation cannot pass on the DOCX structure alone.
    """

    import pymupdf

    preview_root.mkdir(parents=True, exist_ok=True)
    document = pymupdf.open(pdf_path)
    try:
        assert len(document) >= 7
        page_texts = [page.get_text() for page in document]
        combined_text = "\n".join(page_texts)
        for expected in (
            "Тестовый университет",
            "СОДЕРЖАНИЕ",
            "ВВЕДЕНИЕ",
            "ПРИЛОЖЕНИЕ А",
            "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
            "Таблица 1",
            "Широкая таблица Office matrix",
            "Колонка 8",
            "R17C8",
            "Рисунок 1",
            "Контрольное изображение",
            "√",
            "(1)",
        ):
            assert expected in combined_text
        assert "\ufffd" not in combined_text

        toc_text = next(text for text in page_texts if "СОДЕРЖАНИЕ" in text)
        assert "ВВЕДЕНИЕ" in toc_text and "ПРИЛОЖЕНИЕ А" in toc_text
        table_page_index = next(
            index for index, text in enumerate(page_texts) if "Широкая таблица Office matrix" in text
        )
        assert document[table_page_index].rect.width > document[table_page_index].rect.height

        cyrillic_fonts: set[str] = set()
        preview_paths: list[str] = []
        for index, page in enumerate(document, start=1):
            for block in page.get_text("dict").get("blocks", []):
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        span_text = str(span.get("text", ""))
                        if any("А" <= char <= "я" or char in "Ёё" for char in span_text):
                            font_name = str(span.get("font", "")).strip()
                            if font_name:
                                cyrillic_fonts.add(font_name)

            text_blocks = [block for block in page.get_text("blocks") if block[4].strip()]
            assert text_blocks, f"PDF page {index} has no visible text blocks"
            for x0, y0, x1, y1, *_ in text_blocks:
                assert -1 <= x0 <= x1 <= page.rect.width + 1
                assert -1 <= y0 <= y1 <= page.rect.height + 1

            preview_path = preview_root / f"page-{index:04d}.png"
            page.get_pixmap(matrix=pymupdf.Matrix(1.5, 1.5), alpha=False).save(preview_path)
            preview_paths.append(str(preview_path))
            with Image.open(preview_path) as preview:
                assert preview.width > 600 and preview.height > 600
                assert (preview.width > preview.height) is (page.rect.width > page.rect.height)
                assert ImageStat.Stat(preview.convert("L")).extrema[0][0] < 128

            if index > 1:
                footer_text = " ".join(
                    block[4] for block in text_blocks if block[1] > page.rect.height * 0.85
                )
                assert str(index) in footer_text

        assert cyrillic_fonts, "The rendered PDF has no named font span containing Cyrillic text"
        return {
            "pages": len(document),
            "landscape_page": table_page_index + 1,
            "cyrillic_fonts": sorted(cyrillic_fonts),
            "page_previews": preview_paths,
        }
    finally:
        document.close()


@pytest.mark.skipif(
    os.getenv("PAPERCRAFT_RUN_OFFICE_TESTS") != "1",
    reason="set PAPERCRAFT_RUN_OFFICE_TESTS=1 to exercise installed LibreOffice",
)
def test_real_libreoffice_feature_matrix(tmp_path: Path) -> None:
    output_root = Path(os.getenv("PAPERCRAFT_OFFICE_OUTPUT_DIR", str(tmp_path))).resolve()
    output_root.mkdir(parents=True, exist_ok=True)
    image = output_root / "office-matrix-figure.png"
    canvas = Image.new("RGB", (1200, 650), "white")
    drawing = ImageDraw.Draw(canvas)
    drawing.rectangle((80, 80, 1120, 570), outline="black", width=8)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    font = ImageFont.truetype(str(font_path), 62) if font_path.is_file() else None
    drawing.text((230, 280), "PaperCraft Office Matrix", fill="black", font=font)
    canvas.save(image)
    manuscript = Manuscript(
        project_id="office-feature-matrix",
        title="Проверка Office matrix",
        metadata={
            "title_page": {
                "university": "Тестовый университет",
                "department": "Кафедра информационных систем",
                "work_type": "ОТЧЁТ",
                "student": "Студент А. А.",
                "supervisor": "Руководитель Б. Б.",
                "city": "Москва",
                "year": 2026,
            }
        },
        blocks=[
            HeadingBlock(text="ВВЕДЕНИЕ", level=1),
            ParagraphBlock(
                text="Документ проверяет поля, секции, изображения, формулы и приложения. " * 20
            ),
            TableBlock(
                spec=TableSpec(
                    caption="Широкая таблица Office matrix",
                    headers=[f"Колонка {index}" for index in range(1, 9)],
                    rows=[[f"R{row}C{column}" for column in range(1, 9)] for row in range(1, 18)],
                )
            ),
            FormulaBlock(spec=FormulaSpec(expression="x = (-b ± √(b²-4ac))/(2a)", label="1")),
            FigureBlock(caption="Контрольное изображение", artifact_id="matrix-figure"),
            AppendixBlock(
                title="ПРИЛОЖЕНИЕ А",
                blocks=[ParagraphBlock(text="Материалы приложения и контроль нумерации страниц.")],
            ),
        ],
        bibliography=[
            BibliographyEntry(
                title="Официальная документация",
                authors=["A. Author"],
                year=2026,
                publisher="Publisher",
                url="https://example.org/official",
                accessed_on=date(2026, 8, 21),
            )
        ],
    )
    docx = output_root / "office-feature-matrix.docx"
    DocxRenderer(RenderConfig()).render(
        manuscript,
        docx,
        artifact_paths={"matrix-figure": image},
    )
    _append_ref_field(docx, "pc_Table_1")

    with zipfile.ZipFile(docx) as archive:
        document_xml = archive.read("word/document.xml")
        field_parts = [
            archive.read(name)
            for name in archive.namelist()
            if name.startswith(("word/header", "word/footer")) and name.endswith(".xml")
        ]
    assert b"TOC " in document_xml
    assert b"SEQ Table" in document_xml and b"SEQ Figure" in document_xml
    assert b"REF pc_Table_1" in document_xml
    assert b"bookmarkStart" in document_xml
    assert b"oMath" in document_xml
    assert any(b"PAGE" in item for item in field_parts)
    loaded = Document(docx)
    assert len(loaded.sections) >= 3
    landscape_sections = [section for section in loaded.sections if section.orientation == 1]
    assert landscape_sections
    assert all(section.page_width > section.page_height for section in landscape_sections)
    assert loaded.sections[-1].page_height > loaded.sections[-1].page_width
    assert loaded.inline_shapes
    assert "ПРИЛОЖЕНИЕ А" in "\n".join(item.text for item in loaded.paragraphs)

    finalizer = DocumentFinalizer(timeout_seconds=120)
    matrix: dict[str, object] = {
        "docx": str(docx),
        "features": {
            "toc": True,
            "page": True,
            "seq": True,
            "ref": True,
            "bookmarks": True,
            "headers_footers": True,
            "title_page": True,
            "landscape_table": True,
            "formula": True,
            "image": True,
            "bibliography": True,
            "appendix": True,
        },
    }
    if not finalizer.libreoffice_available():
        pytest.skip("LibreOffice is not available")
    libreoffice_pdf = output_root / "office-feature-matrix-libreoffice.pdf"
    libreoffice_result = finalizer.finalize(
        docx,
        pdf_path=libreoffice_pdf,
        preferred="libreoffice",
    )
    assert libreoffice_result.engine == "libreoffice"
    assert libreoffice_result.pdf is not None and libreoffice_result.pdf.valid_header
    visual_qa = _assert_rendered_pdf_feature_matrix(
        libreoffice_pdf,
        output_root / "office-matrix-pages",
    )
    matrix["libreoffice"] = {
        "status": "passed",
        "pdf": str(libreoffice_pdf),
        "size_bytes": libreoffice_result.pdf.size_bytes,
        "visual_qa": visual_qa,
    }
    (output_root / "office-matrix.json").write_text(
        json.dumps(matrix, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
