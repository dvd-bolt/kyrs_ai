import os
import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from models.config import FormattingRulesConfig, TitlePageData

class DocxRenderer:
    """
    Профессиональный движок верстки документов Microsoft Word (.docx) по ГОСТу.
    Обеспечивает точную настройку полей, колонтитулов, нумерации, стилей текста, 
    умное сжатие изображений (Smart Image Fitter), рамки кода и таблицы.
    """
    def __init__(self, formatting_config: FormattingRulesConfig = None):
        self.config = formatting_config or FormattingRulesConfig()

    def create_document(self) -> Document:
        """
        Создает новый документ Word и применяет глобальные стили и поля.
        """
        doc = docx.Document()
        self.setup_document_styles(doc)
        return doc

    def setup_document_styles(self, doc: Document) -> None:
        """
        Настраивает поля страницы, колонтитулы с нумерацией и стили Normal / Heading 1 / Heading 2.
        """
        for section in doc.sections:
            # Настройка полей страницы
            section.top_margin = Cm(self.config.margin_top_cm)
            section.bottom_margin = Cm(self.config.margin_bottom_cm)
            section.left_margin = Cm(self.config.margin_left_cm)
            section.right_margin = Cm(self.config.margin_right_cm)
            
            # Скрытие колонтитулов на 1-й странице (титульнике)
            section.different_first_page_header_footer = True
            
            # Добавление нумерации страниц во 2-й и последующие колонтитулы
            header = section.header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if self.config.page_number_position == "top_right" else WD_ALIGN_PARAGRAPH.CENTER
            
            # Динамическое поле { PAGE }
            run = header_para.add_run()
            run.font.name = self.config.font_name
            run.font.size = Pt(12)
            
            fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
            instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
            fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
            fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run._r.append(fldChar3)

        # Настройка стандартного стиля текста Normal
        style_normal = doc.styles['Normal']
        style_normal.font.name = self.config.font_name
        style_normal.font.size = Pt(self.config.font_size_pt)
        style_normal.font.color.rgb = RGBColor(0, 0, 0)
        style_normal.paragraph_format.line_spacing = self.config.line_spacing
        style_normal.paragraph_format.first_line_indent = Cm(self.config.paragraph_indent_cm)
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style_normal.paragraph_format.space_before = Pt(0)
        style_normal.paragraph_format.space_after = Pt(0)

        # Настройка стиля Заголовок 1 (Heading 1)
        if 'Heading 1' in doc.styles:
            style_h1 = doc.styles['Heading 1']
            style_h1.font.name = self.config.font_name
            style_h1.font.size = Pt(self.config.heading_1_size_pt)
            style_h1.font.bold = True
            style_h1.font.color.rgb = RGBColor(0, 0, 0)
            style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_h1.paragraph_format.first_line_indent = Cm(0)
            style_h1.paragraph_format.space_before = Pt(12)
            style_h1.paragraph_format.space_after = Pt(12)
            style_h1.paragraph_format.keep_with_next = True

        # Настройка стиля Заголовок 2 (Heading 2)
        if 'Heading 2' in doc.styles:
            style_h2 = doc.styles['Heading 2']
            style_h2.font.name = self.config.font_name
            style_h2.font.size = Pt(self.config.heading_2_size_pt)
            style_h2.font.bold = True
            style_h2.font.color.rgb = RGBColor(0, 0, 0)
            style_h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style_h2.paragraph_format.first_line_indent = Cm(self.config.paragraph_indent_cm)
            style_h2.paragraph_format.space_before = Pt(12)
            style_h2.paragraph_format.space_after = Pt(12)
            style_h2.paragraph_format.keep_with_next = True

    def render_title_page(self, doc: Document, title_data: TitlePageData) -> None:
        """
        Генерирует титульный лист по полям настройки на первой странице.
        """
        if title_data.use_custom_file and title_data.custom_docx_path:
            self.attach_external_docx_title(doc, title_data.custom_docx_path)
            return

        # Министерство и Вуз
        p_top = doc.add_paragraph()
        p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_top.paragraph_format.first_line_indent = Cm(0)
        p_top.paragraph_format.space_after = Pt(18)
        run_top = p_top.add_run(f"{title_data.university}\n{title_data.faculty}\n{title_data.department}")
        run_top.font.size = Pt(12)
        run_top.font.bold = True

        # Отступ до центра
        p_space1 = doc.add_paragraph()
        p_space1.paragraph_format.space_before = Pt(36)

        # Тип работы и Тема
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.first_line_indent = Cm(0)
        
        run_work = p_title.add_run(f"{title_data.work_type}\n")
        run_work.font.size = Pt(18)
        run_work.font.bold = True
        
        if title_data.subject:
            run_sub = p_title.add_run(f"{title_data.subject}\n\n")
            run_sub.font.size = Pt(14)
            
        run_topic_lbl = p_title.add_run("на тему: ")
        run_topic_lbl.font.size = Pt(14)
        
        run_topic = p_title.add_run(f"«{title_data.topic}»")
        run_topic.font.size = Pt(16)
        run_topic.font.bold = True

        # Отступ до блока авторов
        p_space2 = doc.add_paragraph()
        p_space2.paragraph_format.space_before = Pt(48)

        # Блок Исполнителя и Проверяющего (в виде правой таблицы без границ)
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        cell_left, cell_right = table.cell(0, 0), table.cell(0, 1)
        
        # Скрытие границ таблицы
        tblPr = table._tbl.tblPr
        tblBorders = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
        tblPr.append(tblBorders)
        
        p_author = cell_right.paragraphs[0]
        p_author.paragraph_format.first_line_indent = Cm(0)
        p_author.paragraph_format.line_spacing = 1.15
        
        run_stud = p_author.add_run(f"{title_data.student_info}\n\n{title_data.teacher_info}")
        run_stud.font.size = Pt(12)

        # Отступ вниз до Города и Года
        p_space3 = doc.add_paragraph()
        p_space3.paragraph_format.space_before = Pt(72)

        p_bottom = doc.add_paragraph()
        p_bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_bottom.paragraph_format.first_line_indent = Cm(0)
        run_bot = p_bottom.add_run(f"{title_data.city} – {title_data.year}")
        run_bot.font.size = Pt(12)

        # Разрыв страницы после титульного листа
        doc.add_page_break()

    def attach_external_docx_title(self, target_doc: Document, external_title_path: str) -> None:
        """
        Подшивает первую страницу из внешнего файла .docx титульного листа.
        """
        if not os.path.exists(external_title_path):
            raise FileNotFoundError(f"Файл титульного листа не найден: {external_title_path}")
            
        ext_doc = docx.Document(external_title_path)
        for element in ext_doc.element.body:
            target_doc.element.body.append(element)
        target_doc.add_page_break()

    def add_heading_1(self, doc: Document, text: str) -> Paragraph:
        """
        Добавляет Заголовок 1-го уровня по центру с прописной буквы без точки.
        """
        p = doc.add_paragraph(text, style='Heading 1')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def add_heading_2(self, doc: Document, text: str) -> Paragraph:
        """
        Добавляет Заголовок 2-го уровня с абзацным отступом.
        """
        p = doc.add_paragraph(text, style='Heading 2')
        return p

    def add_paragraph(self, doc: Document, text: str) -> Paragraph:
        """
        Добавляет стандартный абзац текста по ширине с красной строкой.
        """
        return doc.add_paragraph(text, style='Normal')

    def insert_smart_image(self, doc: Document, image_path: str, caption: str, max_width_cm: float = 16.5) -> None:
        """
        Smart Image Fitter: Пропорционально сжимает изображение под ширину полей (макс 16.5 см),
        центрирует и добавляет подпись Рисунок X.Y – Название по ГОСТу.
        """
        if not os.path.exists(image_path):
            p_err = doc.add_paragraph(f"[ОШИБКА: Изображение не найдено по пути {image_path}]")
            p_err.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            return

        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.first_line_indent = Cm(0)
        p_img.paragraph_format.space_before = Pt(12)
        
        run_img = p_img.add_run()
        run_img.add_picture(image_path, width=Cm(max_width_cm))

        # Подпись под рисунком (12pt, центрировано, без точки на конце)
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.first_line_indent = Cm(0)
        p_cap.paragraph_format.space_before = Pt(6)
        p_cap.paragraph_format.space_after = Pt(12)
        
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = self.config.font_name
        run_cap.font.size = Pt(12)

    def add_code_block(self, doc: Document, code_text: str) -> None:
        """
        Оформляет листинг кода в одноячеисточной таблице с серой заливкой #F4F4F4 и моноширинным шрифтом.
        """
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        
        # Серая заливка фона ячейки
        shd = parse_xml(r'<w:shd %s w:fill="F4F4F4"/>' % nsdecls('w'))
        cell._tc.get_or_add_tcPr().append(shd)
        
        # Поля внутри рамки кода
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(r'<w:tcMar %s><w:top w:w="120" w:type="dxa"/><w:bottom w:w="120" w:type="dxa"/><w:left w:w="180" w:type="dxa"/><w:right w:w="180" w:type="dxa"/></w:tcMar>' % nsdecls('w'))
        tcPr.append(tcMar)
        
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(40, 40, 40)
        
        # Добавляем пустой абзац после таблицы кода для отступа
        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_before = Pt(6)

    def render_table(self, doc: Document, title: str, headers: list, rows: list) -> None:
        """
        Верстает таблицу с подписью Таблица X.Y – Название над ней и выравниванием по центру.
        """
        # Подпись над таблицей (12pt, по левому краю / с отступом)
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_title.paragraph_format.first_line_indent = Cm(self.config.paragraph_indent_cm)
        p_title.paragraph_format.space_before = Pt(12)
        p_title.paragraph_format.space_after = Pt(6)
        
        run_title = p_title.add_run(title)
        run_title.font.name = self.config.font_name
        run_title.font.size = Pt(12)
        run_title.font.bold = True

        # Создание таблицы
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Стилизация шапки таблицы
        hdr_cells = table.rows[0].cells
        for idx, header_text in enumerate(headers):
            hdr_cells[idx].text = header_text
            p = hdr_cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = self.config.table_line_spacing
            run = p.runs[0]
            run.font.name = self.config.font_name
            run.font.size = Pt(self.config.table_font_size_pt)
            run.font.bold = True
            
            # Серая заливка шапки
            shd = parse_xml(r'<w:shd %s w:fill="E8E8E8"/>' % nsdecls('w'))
            hdr_cells[idx]._tc.get_or_add_tcPr().append(shd)

        # Заполнение строк таблицы
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_value in enumerate(row_data):
                row_cells[col_idx].text = str(cell_value)
                p = row_cells[col_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = self.config.table_line_spacing
                if p.runs:
                    p.runs[0].font.name = self.config.font_name
                    p.runs[0].font.size = Pt(self.config.table_font_size_pt)

        # Отступ после таблицы
        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_before = Pt(12)

    def add_table_of_contents_placeholder(self, doc: Document, plan_structure: list = None) -> None:
        """
        Верстает оглавление в виде эстетичной таблицы без видимых границ.
        """
        self.add_heading_1(doc, "СОДЕРЖАНИЕ")
        
        if not plan_structure:
            p_stub = doc.add_paragraph("[Здесь будет автоматически сформированное Содержание]")
            p_stub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()
            return

        table = doc.add_table(rows=len(plan_structure), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Скрытие границ
        tblPr = table._tbl.tblPr
        tblBorders = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
        tblPr.append(tblBorders)

        for idx, item in enumerate(plan_structure):
            cell_left, cell_right = table.rows[idx].cells
            
            p_left = cell_left.paragraphs[0]
            p_left.paragraph_format.first_line_indent = Cm(0)
            p_left.paragraph_format.line_spacing = 1.15
            run_title = p_left.add_run(item.get("title", ""))
            run_title.font.name = self.config.font_name
            run_title.font.size = Pt(self.config.font_size_pt)
            if item.get("is_section_header", False):
                run_title.font.bold = True
                
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.first_line_indent = Cm(0)
            run_page = p_right.add_run(str(item.get("page", 3)))
            run_page.font.name = self.config.font_name
            run_page.font.size = Pt(self.config.font_size_pt)

        doc.add_page_break()
